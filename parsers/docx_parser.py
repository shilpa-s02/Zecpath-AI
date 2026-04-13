import docx
from utils.logger import log

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extracts text from a Word document (.docx).
    1. Iterates through all paragraphs in order.
    2. Iterates through tables and formats them as readable text.
    3. Merges them into a single string.
    """
    if not docx_path:
        return ""

    full_text = []

    try:
        doc = docx.Document(docx_path)
        
        # 1. Extract block-level elements in order
        # We can iterate through the document's body element
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                para = [p for p in doc.paragraphs if p._element == element]
                if para:
                    full_text.append(para[0].text)
            
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                table = [t for t in doc.tables if t._element == element]
                if table:
                    table_text = ""
                    for row in table[0].rows:
                        # Extract text from each cell and join with |
                        clean_row = [cell.text.strip() for cell in row.cells]
                        table_text += " | ".join(clean_row) + "\n"
                    full_text.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")

        return "\n".join(full_text)

    except Exception as e:
        log.error(f"Error parsing DOCX {docx_path}: {e}")
        return ""
